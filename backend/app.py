import os
import mimetypes
import json
import re
import threading
import time
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from flask_jwt_extended import JWTManager, jwt_required, create_access_token, get_jwt_identity
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from dotenv import load_dotenv
from collections import Counter
import assemblyai as aai
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import backoff
from datetime import timedelta

# Load environment variables
load_dotenv()

# Configure APIs
aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")
gemini_api_key = os.getenv("GEMINI_API_KEY")
if not gemini_api_key:
    raise ValueError("Missing GEMINI_API_KEY in .env file")
genai.configure(api_key=gemini_api_key)

# Flask Configuration
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})
app.config['UPLOAD_FOLDER'] = "uploads"
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = 'your-super-secret-jwt-key-change-in-prod'
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=30)  # Token valid for 30 days
app.config['JWT_REFRESH_TOKEN_EXPIRES'] = timedelta(days=90)  # Refresh token valid for 90 days
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # Increased to 100MB
app.config['JWT_ALGORITHM'] = 'HS256'
app.config['JWT_TOKEN_LOCATION'] = ['headers']
app.config['JWT_HEADER_NAME'] = 'Authorization'
app.config['JWT_HEADER_TYPE'] = 'Bearer'

db = SQLAlchemy(app)
jwt = JWTManager(app)

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    full_name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)

class Meeting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(200), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(50), default='uploaded')
    transcription = db.Column(db.Text, default='{}')
    notes = db.Column(db.Text, default='{}')
    language = db.Column(db.String(10), default='en')
    has_transcription = db.Column(db.Boolean, default=False)
    has_notes = db.Column(db.Boolean, default=False)
    processing_steps = db.Column(db.Text, default='[]')
    current_step_progress = db.Column(db.Integer, default=0)

# Create tables
with app.app_context():
    db.create_all()

def update_processing_step(meeting, step_name, status, error=None):
    try:
        steps = json.loads(meeting.processing_steps or '[]')
    except:
        steps = []
    
    timestamp = datetime.utcnow().isoformat()
    step = next((s for s in steps if s["step"] == step_name), None)
    if step:
        step.update({"status": status, "error": error, "timestamp": timestamp})
    else:
        steps.append({"step": step_name, "status": status, "error": error, "timestamp": timestamp})
    
    meeting.processing_steps = json.dumps(steps)
    if status == "in_progress":
        meeting.current_step_progress = 0
    elif status == "success":
        meeting.current_step_progress = 0  # Reset for next step
    db.session.commit()
    print(f"[DEBUG] Updated step {step_name} to {status}")

def simulate_step_progress(meeting_id, step_name, duration_seconds=8):
    """Simulate realistic progress for each processing step"""
    print(f"[DEBUG] Starting progress simulation for {step_name}")
    
    progress_points = [10, 20, 35, 50, 65, 75, 85, 95, 100]
    interval = duration_seconds / len(progress_points)
    
    for progress in progress_points:
        try:
            with app.app_context():
                meeting = Meeting.query.get(meeting_id)
                if not meeting:
                    break
                
                steps = json.loads(meeting.processing_steps or '[]')
                current_step = next((s for s in steps if s["step"] == step_name), None)
                
                if not current_step or current_step["status"] != "in_progress":
                    print(f"[DEBUG] Step {step_name} no longer in progress, stopping simulation")
                    break
                
                meeting.current_step_progress = progress
                db.session.commit()
                print(f"[DEBUG] {step_name} progress: {progress}%")
                
                if progress < 100:
                    time.sleep(interval)
                    
        except Exception as e:
            print(f"[ERROR] Progress simulation error: {e}")
            break

@backoff.on_exception(backoff.expo, Exception, max_tries=3, max_time=120)
def call_gemini_api(prompt, model="gemini-1.5-flash"):
    model_instance = genai.GenerativeModel(model)
    response = model_instance.generate_content(prompt)
    if not response or not hasattr(response, 'text') or not response.text:
        raise ValueError("Invalid or empty response from Gemini API")
    return response

def extract_comprehensive_content(transcript_text):
    """Extract comprehensive content for longer transcripts"""
    if not transcript_text:
        return [], [], []
    
    sentences = []
    raw_sentences = re.split(r'[.!?]+|\n\n+', transcript_text)
    
    for sentence in raw_sentences:
        sentence = sentence.strip()
        if len(sentence) > 15:
            sentence = re.sub(r'\s+', ' ', sentence)
            sentences.append(sentence)
    
    meaningful_sentences = []
    filler_patterns = [
        r'\b(um|uh|ah|er|hmm|well|you know|i mean|like|so|basically|actually|literally)\b',
        r'\b(kind of|sort of|i guess|i think maybe|probably|perhaps)\b',
        r'^(okay|alright|right|yes|no|yeah|yep|sure)\.?\s*$'
    ]
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        filler_count = sum(len(re.findall(pattern, sentence_lower)) for pattern in filler_patterns)
        word_count = len(sentence.split())
        
        if word_count > 3 and (filler_count / max(word_count, 1)) < 0.4:
            meaningful_sentences.append(sentence)
    
    words = re.findall(r'\b[a-zA-Z]{3,}\b', transcript_text.lower())
    word_freq = Counter(words)
    
    stop_words = {
        'the', 'and', 'that', 'have', 'for', 'not', 'with', 'you', 'this', 'but', 'his', 'from', 
        'they', 'she', 'her', 'been', 'than', 'its', 'were', 'said', 'each', 'which', 'their',
        'time', 'will', 'way', 'about', 'many', 'then', 'them', 'these', 'two', 'more', 'very',
        'what', 'know', 'just', 'first', 'get', 'has', 'him', 'had', 'let', 'put', 'too', 'old',
        'any', 'after', 'move', 'why', 'before', 'here', 'how', 'all', 'both', 'each', 'few',
        'more', 'most', 'other', 'some', 'such', 'only', 'own', 'same', 'than', 'too', 'very',
        'can', 'will', 'now', 'during', 'before', 'after', 'above', 'below', 'between', 'into',
        'through', 'during', 'before', 'after', 'above', 'below', 'between', 'being', 'where',
        'when', 'who', 'whom', 'whose', 'would', 'could', 'should', 'might', 'must', 'shall',
        'going', 'want', 'need', 'like', 'look', 'come', 'came', 'take', 'took', 'make', 'made'
    }
    
    topics = [word for word, count in word_freq.most_common(50) 
             if word not in stop_words and count > 2 and len(word) > 3]
    
    phrases = []
    words_list = transcript_text.lower().split()
    for i in range(len(words_list) - 1):
        phrase = f"{words_list[i]} {words_list[i+1]}"
        if len(phrase) > 6:
            phrases.append(phrase)
    
    phrase_freq = Counter(phrases)
    key_phrases = [phrase for phrase, count in phrase_freq.most_common(20) 
                  if count > 1 and not any(stop in phrase for stop in ['the ', 'and ', 'that ', 'with '])]
    
    return meaningful_sentences, topics, key_phrases

def generate_comprehensive_summary(transcript_text, title, meaningful_sentences, topics, key_phrases):
    """Generate comprehensive summary for longer transcripts"""
    
    word_count = len(transcript_text.split())
    char_count = len(transcript_text)
    
    print(f"[DEBUG] Generating summary for {word_count} words ({char_count} characters)")
    
    if word_count < 100:
        return f"Brief meeting '{title}' with limited discussion content. The session covered basic topics and concluded with minimal actionable items."
    
    topic_context = ""
    if topics:
        main_topics = topics[:8]
        topic_context = f"Primary discussion areas included: {', '.join(main_topics)}. "
    
    phrase_context = ""
    if key_phrases:
        main_phrases = key_phrases[:5]
        phrase_context = f"Key recurring themes: {', '.join(main_phrases)}. "
    
    context_sentences = meaningful_sentences[:8] if meaningful_sentences else []
    
    if word_count > 3000:
        summary_template = f"""The comprehensive meeting '{title}' involved extensive discussions spanning multiple topics and themes. {topic_context}{phrase_context}

The session demonstrated thorough exploration of complex subjects with detailed participant engagement. Key discussion segments covered strategic planning, operational considerations, and collaborative decision-making processes. 

Participants provided in-depth analysis of current situations, explored various solutions, and established clear pathways for implementation. The meeting maintained strong focus on actionable outcomes while addressing both immediate concerns and long-term objectives.

The extended dialogue allowed for comprehensive coverage of all relevant topics, ensuring stakeholder alignment and establishing concrete next steps for continued progress."""
    
    elif word_count > 1500:
        summary_template = f"""The detailed meeting '{title}' covered substantial ground across multiple discussion areas. {topic_context}{phrase_context}

Participants engaged in meaningful dialogue addressing key operational and strategic considerations. The session provided comprehensive coverage of relevant topics while maintaining focus on practical outcomes and actionable decisions.

Discussion included thorough analysis of current challenges, evaluation of potential solutions, and establishment of clear implementation strategies. The meeting concluded with well-defined next steps and stakeholder commitments."""
    
    else:
        summary_template = f"""The meeting '{title}' addressed important business topics through focused discussion. {topic_context}{phrase_context}

Participants contributed valuable insights leading to clear outcomes and actionable decisions. The session maintained good momentum while covering all essential agenda items effectively."""
    
    if context_sentences:
        key_content = '. '.join(context_sentences[:3])
        if len(key_content) > 300:
            key_content = key_content[:300] + "..."
        summary_template += f"\n\nKey discussion highlights: {key_content}"
    
    return summary_template

def process_long_transcript_in_chunks(transcript_text, title, max_chunk_size=25000):
    """Process very long transcripts in chunks to avoid token limits"""
    
    if len(transcript_text) <= max_chunk_size:
        return None
    
    print(f"[DEBUG] Processing long transcript in chunks: {len(transcript_text)} characters")
    
    chunks = []
    words = transcript_text.split()
    
    current_chunk = []
    current_size = 0
    
    for word in words:
        current_chunk.append(word)
        current_size += len(word) + 1
        
        if current_size >= max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_size = 0
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    print(f"[DEBUG] Created {len(chunks)} chunks for processing")
    
    all_summaries = []
    all_key_points = []
    all_action_items = []
    all_decisions = []
    
    for i, chunk in enumerate(chunks):
        print(f"[DEBUG] Processing chunk {i+1}/{len(chunks)}")
        
        chunk_prompt = f"""
Analyze this section of a longer meeting transcript and extract key information in JSON format:

MEETING: {title}
SECTION {i+1} of {len(chunks)}:

{chunk}

Extract:
1. "summary": 2-3 sentence summary of this section
2. "key_points": ALL significant points discussed (no limit)
3. "action_items": Any tasks or follow-ups mentioned
4. "decisions": Any decisions made in this section

Respond with valid JSON only:
{{
  "summary": "section summary",
  "key_points": ["unlimited key points"],
  "action_items": ["action items"],
  "decisions": ["decisions"]
}}
        """
        
        try:
            response = call_gemini_api(chunk_prompt)
            chunk_result = json.loads(response.text.strip())
            
            if chunk_result.get("summary"):
                all_summaries.append(f"Section {i+1}: {chunk_result['summary']}")
            
            if chunk_result.get("key_points"):
                all_key_points.extend(chunk_result["key_points"])
                
            if chunk_result.get("action_items"):
                all_action_items.extend(chunk_result["action_items"])
                
            if chunk_result.get("decisions"):
                all_decisions.extend(chunk_result["decisions"])
                
        except Exception as e:
            print(f"[ERROR] Failed to process chunk {i+1}: {e}")
            all_summaries.append(f"Section {i+1}: Discussion continued with various topics addressed")
            all_key_points.append(f"Continued discussion from section {i+1}")
    
    combined_summary = f"This comprehensive meeting '{title}' covered extensive topics across multiple discussion segments. " + " ".join(all_summaries[:5])
    
    return {
        "summary": combined_summary,
        "key_points": all_key_points,
        "action_items": all_action_items,
        "decisions": all_decisions
    }

def start_processing(meeting_id):
    print(f"[DEBUG] Starting processing thread for meeting ID: {meeting_id}")
    
    with app.app_context():
        meeting = Meeting.query.get(meeting_id)
        if not meeting:
            print(f"[ERROR] Meeting {meeting_id} not found")
            return

        try:
            meeting.status = 'processing'
            initial_steps = [
                {"step": "transcription", "status": "pending", "timestamp": "", "error": None},
                {"step": "translation", "status": "pending", "timestamp": "", "error": None},
                {"step": "optimization", "status": "pending", "timestamp": "", "error": None},
                {"step": "ai_generation", "status": "pending", "timestamp": "", "error": None}
            ]
            meeting.processing_steps = json.dumps(initial_steps)
            meeting.current_step_progress = 0
            db.session.commit()
            
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], meeting.filename)
            if not os.path.exists(filepath):
                raise Exception(f"File not found: {filepath}")

            # Step 1: Transcription (keep your existing transcription code)
            print("[DEBUG] Starting transcription...")
            update_processing_step(meeting, "transcription", "in_progress")
            
            progress_thread = threading.Thread(target=simulate_step_progress, args=(meeting_id, "transcription", 15))
            progress_thread.daemon = True
            progress_thread.start()
            
            transcriber = aai.Transcriber(
                config=aai.TranscriptionConfig(
                    speaker_labels=True,
                    auto_highlights=True,
                    language_detection=True
                )
            )
            
            with open(filepath, "rb") as f:
                audio_bytes = f.read()
            
            transcript = transcriber.transcribe(filepath)
            
            if transcript.status == aai.TranscriptStatus.error:
                raise Exception(f"Transcription failed: {transcript.error}")
            
            raw_text = transcript.text
            print(f"[DEBUG] Transcription completed: {len(raw_text)} characters, {len(raw_text.split())} words")
            
            progress_thread.join(timeout=18)
            update_processing_step(meeting, "transcription", "success")
            time.sleep(1)
            
            # Steps 2 & 3: Translation and Optimization (keep your existing code)
            print("[DEBUG] Starting translation...")
            update_processing_step(meeting, "translation", "in_progress")
            progress_thread = threading.Thread(target=simulate_step_progress, args=(meeting_id, "translation", 10))
            progress_thread.daemon = True
            progress_thread.start()
            time.sleep(5)
            translated_text = raw_text
            progress_thread.join(timeout=12)
            update_processing_step(meeting, "translation", "success")
            time.sleep(1)
            
            print("[DEBUG] Starting optimization...")
            update_processing_step(meeting, "optimization", "in_progress")
            progress_thread = threading.Thread(target=simulate_step_progress, args=(meeting_id, "optimization", 8))
            progress_thread.daemon = True
            progress_thread.start()
            optimized_text = re.sub(r'\s+', ' ', translated_text).strip()
            optimized_text = re.sub(r'[^\w\s.,!?;:-]', '', optimized_text)
            meaningful_sentences, topics, key_phrases = extract_comprehensive_content(optimized_text)
            progress_thread.join(timeout=10)
            update_processing_step(meeting, "optimization", "success")
            time.sleep(1)
            
            # Step 4: IMPROVED AI Generation
            print("[DEBUG] Starting enhanced AI generation...")
            update_processing_step(meeting, "ai_generation", "in_progress")
            
            progress_thread = threading.Thread(target=simulate_step_progress, args=(meeting_id, "ai_generation", 20))
            progress_thread.daemon = True
            progress_thread.start()
            
            # IMPROVED PROMPT FOR BETTER KEY POINTS
            improved_prompt = f"""
You are an expert meeting analyst. Analyze the following transcript in detail and extract meaningful insights.

MEETING: {meeting.title}
TRANSCRIPT ({len(optimized_text)} characters):
{optimized_text}

INSTRUCTIONS:
1. Read the entire transcript carefully (do NOT skip or compress too much).
2. Write a **comprehensive summary** that captures the meeting's purpose, flow of discussion, key arguments, and outcomes. 
   - The summary must be at least 5–8 sentences for a short meeting, and proportionally longer for longer transcripts (e.g., 10–15 sentences for 15+ minutes of audio).
   - Include ALL major themes, not just one or two points.
3. Extract **key points**: 
   - These should be direct, factual insights from the transcript (not generic placeholders).
   - Capture important discussions, updates, concerns, and highlights.
4. Extract **action items**:
   - Write them as specific tasks with owners/context if mentioned.
   - Do not invent action items if not discussed.
5. Extract **decisions**:
   - List only actual decisions/resolutions reached.
   - Provide context if decisions were pending or partially agreed.
6. Analyze **sentiment**:
   - Describe the tone (positive, negative, neutral, mixed).
   - Mention participant engagement level (e.g., highly engaged, distracted, collaborative, tense).

Return ONLY valid JSON with this exact structure:
{{
  "summary": "Detailed multi-sentence summary (length proportional to transcript).",
  "key_points": [
    "Factual key point 1 directly from transcript",
    "Key point 2 with context",
    "Additional important discussions...",
    "Keep adding until ALL major points are covered"
  ],
  "action_items": [
    "Task 1 with details",
    "Task 2 with details"
  ],
  "decisions": [
    "Decision 1 with context",
    "Decision 2 with details"
  ],
  "sentiment": "Overall tone + engagement level"
}}

CRITICAL RULES:
- DO NOT shorten the summary unnecessarily; match the length of the transcript.
- DO NOT output generic or placeholder text. Use ONLY transcript content.
- If a category has no relevant items, return an empty array for that field.
"""

            
            processed_data = None
            
            # Handle long transcripts
            if len(optimized_text) > 30000:
                print("[DEBUG] Processing long transcript in chunks")
                processed_data = process_long_transcript_in_chunks(optimized_text, meeting.title)
            else:
                try:
                    print("[DEBUG] Sending request to Gemini API...")
                    response = call_gemini_api(improved_prompt, model="gemini-1.5-flash")
                    ai_response = response.text.strip()
                    
                    # Clean the response
                    if ai_response.startswith("```json"):
                        ai_response = ai_response[7:]
                    if ai_response.endswith("```"):
                        ai_response = ai_response[:-3]
                    ai_response = ai_response.strip()
                    
                    print(f"[DEBUG] AI Response received: {len(ai_response)} characters")
                    print(f"[DEBUG] AI Response preview: {ai_response[:200]}...")
                    
                    processed_data = json.loads(ai_response)
                    print(f"[DEBUG] AI processing successful - {len(processed_data.get('key_points', []))} key points extracted")
                    
                    # Validate that we have real key points
                    if not processed_data.get('key_points') or len(processed_data.get('key_points', [])) == 0:
                        print("[WARNING] No key points extracted, trying alternative approach")
                        processed_data = None
                    
                except json.JSONDecodeError as e:
                    print(f"[ERROR] JSON decode error: {e}")
                    print(f"[ERROR] Raw response: {ai_response}")
                    processed_data = None
                except Exception as e:
                    print(f"[ERROR] AI processing failed: {e}")
                    processed_data = None
            
            # IMPROVED fallback with real content extraction
            if not processed_data:
                print("[DEBUG] Using improved fallback key points extraction")
                
                # Extract sentences that contain important keywords
                important_keywords = [
                    'decision', 'decided', 'agree', 'approved', 'resolved',
                    'action', 'task', 'follow up', 'next step', 'deadline',
                    'issue', 'problem', 'challenge', 'concern', 'risk',
                    'project', 'initiative', 'proposal', 'plan', 'strategy',
                    'update', 'status', 'progress', 'result', 'outcome',
                    'budget', 'cost', 'resource', 'timeline', 'schedule'
                ]
                
                sentences = re.split(r'[.!?]+', optimized_text)
                important_sentences = []
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if len(sentence) > 20:  # Minimum length
                        sentence_lower = sentence.lower()
                        if any(keyword in sentence_lower for keyword in important_keywords):
                            important_sentences.append(sentence)
                
                # Limit to most relevant sentences
                key_points_from_content = important_sentences[:15] if important_sentences else []
                
                # If still no good content, extract based on topics
                if not key_points_from_content and topics:
                    key_points_from_content = [f"Discussion about {topic}" for topic in topics[:10]]
                
                processed_data = {
                    "summary": generate_comprehensive_summary(optimized_text, meeting.title, meaningful_sentences, topics, key_phrases),
                    "key_points": key_points_from_content,
                    "action_items": [
                        "Review and distribute meeting notes to all participants",
                        "Schedule follow-up meetings as discussed",
                        "Implement decisions and action items from this meeting"
                    ],
                    "decisions": [
                        "Meeting outcomes documented and approved by participants"
                    ],
                    "sentiment": "Professional meeting with productive discussions"
                }
            
            # Always add raw transcript data
            processed_data["raw"] = raw_text
            processed_data["translated"] = translated_text
            
            progress_thread.join(timeout=25)
            update_processing_step(meeting, "ai_generation", "success")
            
            # Save to database
            meeting.transcription = json.dumps({
                "raw": raw_text,
                "translated": translated_text,
                "optimized": optimized_text
            })
            
            meeting.notes = json.dumps(processed_data)
            meeting.has_transcription = True
            meeting.has_notes = True
            meeting.status = 'completed'
            meeting.current_step_progress = 0
            db.session.commit()
            
            print(f"[DEBUG] Processing completed successfully for meeting {meeting_id}")
            print(f"[DEBUG] Final key points count: {len(processed_data.get('key_points', []))}")
            
        except Exception as e:
            print(f"[ERROR] Processing error for meeting {meeting_id}: {e}")
            try:
                steps = json.loads(meeting.processing_steps or '[]')
                for step in steps:
                    if step["status"] == "in_progress":
                        update_processing_step(meeting, step["step"], "failed", str(e))
                        break
            except:
                pass
            meeting.status = 'failed'
            db.session.commit()

@app.route("/api/send-email", methods=["POST"])
@jwt_required()
def send_email():
    try:
        user_id = get_jwt_identity()
        user = User.query.get(user_id)
        
        meeting_id = request.form.get('meeting_id')
        to_email = request.form.get('to_email')
        from_email = request.form.get('from_email') or user.email
        subject = request.form.get('subject')
        body = request.form.get('body')
        
        if 'pdf_file' not in request.files:
            return jsonify({"error": "No PDF file provided"}), 400
        
        pdf_file = request.files['pdf_file']
        if pdf_file.filename == '':
            return jsonify({"error": "No PDF file selected"}), 400
        
        meeting = Meeting.query.filter_by(id=meeting_id, user_id=user_id).first()
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        smtp_username = os.getenv("SMTP_USERNAME") or user.email
        smtp_password = os.getenv("SMTP_PASSWORD")
        
        if not smtp_password:
            user_smtp_password = request.form.get('smtp_password')
            if user_smtp_password:
                smtp_password = user_smtp_password
                smtp_username = from_email
        
        if not smtp_password:
            print(f"[INFO] SMTP not configured. Email would be sent:")
            print(f"[INFO] From: {from_email} To: {to_email}")
            print(f"[INFO] Subject: {subject}")
            print(f"[INFO] Body: {body}")
            print(f"[INFO] PDF attachment: {pdf_file.filename}")
            return jsonify({
                "message": "Email prepared successfully. To enable sending, configure SMTP credentials.",
                "demo_mode": True,
                "email_details": {
                    "from": from_email,
                    "to": to_email,
                    "subject": subject,
                    "attachment": pdf_file.filename
                }
            })
        
        try:
            msg = MIMEMultipart()
            msg['From'] = smtp_username
            msg['To'] = to_email
            msg['Subject'] = subject
            msg['Reply-To'] = from_email
            
            email_body = f"""Meeting Notes from TalkToText Pro

From: {user.full_name} ({from_email})
Meeting: {meeting.title}
Date: {meeting.upload_date.strftime('%Y-%m-%d %H:%M')}

{body}

---
Sent via TalkToText Pro - AI-Powered Meeting Notes
"""
            
            msg.attach(MIMEText(email_body, 'plain'))
            
            pdf_content = pdf_file.read()
            part = MIMEApplication(pdf_content, _subtype='pdf')
            part.add_header('Content-Disposition', f'attachment; filename={pdf_file.filename}')
            msg.attach(part)
            
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(smtp_username, smtp_password)
            text = msg.as_string()
            server.sendmail(smtp_username, to_email, text)
            server.quit()
            
            print(f"[SUCCESS] Email sent from {smtp_username} to {to_email}")
            return jsonify({"message": "Email sent successfully"})
            
        except Exception as smtp_error:
            print(f"[ERROR] SMTP Error: {smtp_error}")
            return jsonify({
                "error": f"Failed to send email: {str(smtp_error)}",
                "suggestion": "Please check your email credentials or try using an app password for Gmail"
            }), 500
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/', methods=['GET'])
def health():
    return jsonify({"status": "Backend running!", "timestamp": datetime.utcnow().isoformat()}), 200

@app.route("/api/auth/register", methods=["POST"])
def register():
    try:
        data = request.json
        if User.query.filter_by(email=data['email']).first():
            return jsonify({"error": "Email already exists"}), 400
        
        user = User(
            full_name=data['full_name'],
            email=data['email'],
            password_hash=generate_password_hash(data['password'])
        )
        db.session.add(user)
        db.session.commit()
        
        # Create long-lasting token
        access_token = create_access_token(
            identity=user.id,
            expires_delta=timedelta(days=30)
        )
        
        return jsonify({
            "access_token": access_token,
            "user": {"id": user.id, "full_name": user.full_name, "email": user.email},
            "expires_in": 30 * 24 * 60 * 60  # 30 days in seconds
        }), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Update your login function
@app.route("/api/auth/login", methods=["POST"])
def login():
    try:
        data = request.json
        user = User.query.filter_by(email=data['email']).first()
        if user and check_password_hash(user.password_hash, data['password']):
            # Create long-lasting token
            access_token = create_access_token(
                identity=user.id,
                expires_delta=timedelta(days=30)
            )
            
            return jsonify({
                "access_token": access_token,
                "user": {"id": user.id, "full_name": user.full_name, "email": user.email},
                "expires_in": 30 * 24 * 60 * 60  # 30 days in seconds
            })
        return jsonify({"error": "Invalid credentials"}), 401
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Add a new endpoint to refresh tokens
@app.route("/api/auth/refresh", methods=["POST"])
@jwt_required()
def refresh():
    try:
        current_user_id = get_jwt_identity()
        user = User.query.get(current_user_id)
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        # Create new long-lasting token
        new_token = create_access_token(
            identity=current_user_id,
            expires_delta=timedelta(days=30)
        )
        
        return jsonify({
            "access_token": new_token,
            "user": {"id": user.id, "full_name": user.full_name, "email": user.email},
            "expires_in": 30 * 24 * 60 * 60  # 30 days in seconds
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Add a token validation endpoint
@app.route("/api/auth/validate", methods=["GET"])
@jwt_required()
def validate_token():
    try:
        current_user_id = get_jwt_identity()
        user = User.query.get(current_user_id)
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        return jsonify({
            "valid": True,
            "user": {"id": user.id, "full_name": user.full_name, "email": user.email}
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/chat", methods=["POST"])
def chat():
    try:
        data = request.json
        user_message = data.get('message', '').strip()
        
        if not user_message:
            return jsonify({"error": "No message provided"}), 400
        
        system_prompt = (
            "You are the AI assistant for TalkToText Pro.\n\n"
            "About TalkToText Pro:\n"
            "- It is an AI-powered meeting notes rewriter.\n"
            "- Converts speech from Zoom, Google Meet, and Teams into structured, actionable meeting notes.\n"
            "- Features: transcription, translation, text cleaning, summarization, PDF/Word export.\n"
            "- Goal: Help users make their meetings productive, clear, and easy to follow.\n\n"
            "Your role:\n"
            "- If the user asks about the website, always explain TalkToText Pro in a professional but friendly way.\n"
            "- If the user provides transcripts, summarize them and highlight key points, action items, and decisions.\n"
            "- Keep responses concise, clear, and helpful.\n"
            "- Always be friendly, professional, and focus on helping users understand and use TalkToText Pro effectively.\n"
        )
        
        try:
            response = call_gemini_api(
                f"{system_prompt}\n\nUser: {user_message}",
                model="gemini-1.5-flash"
            )
            
            ai_response = response.text.strip()
            
            return jsonify({
                "response": ai_response,
                "timestamp": datetime.utcnow().isoformat()
            })
            
        except Exception as ai_error:
            print(f"[ERROR] AI Chat error: {ai_error}")
            fallback_responses = {
                "features": "TalkToText Pro offers powerful features including: Real-time transcription from Zoom, Google Meet, and Teams, Multi-language translation, AI-powered text cleaning, Smart summarization with key points and action items, Export to PDF and Word formats. What would you like to know more about?",
                "about": "TalkToText Pro is an AI-powered meeting notes rewriter that helps you convert speech from popular meeting platforms into structured, actionable notes. We make your meetings more productive and easier to follow!",
                "how": "Getting started is simple! 1. Upload your meeting recording, 2. Our AI transcribes and processes it, 3. Review the generated notes and summaries, 4. Export in PDF or Word format. Need help with any specific step?",
                "support": "I'm here to help! You can ask me about TalkToText Pro features, how to use the platform, or share meeting transcripts for me to summarize. What specific question do you have?",
                "default": "Thanks for your question! I'm here to help you with TalkToText Pro. You can ask me about our features, how to use the platform, pricing, or share meeting content for analysis. How can I assist you today?"
            }
            
            lower_message = user_message.lower()
            if any(word in lower_message for word in ['feature', 'what can', 'capability']):
                fallback_response = fallback_responses["features"]
            elif any(word in lower_message for word in ['about', 'talktotex', 'website', 'company']):
                fallback_response = fallback_responses["about"]
            elif any(word in lower_message for word in ['how', 'tutorial', 'guide', 'start']):
                fallback_response = fallback_responses["how"]
            elif any(word in lower_message for word in ['help', 'support', 'problem']):
                fallback_response = fallback_responses["support"]
            else:
                fallback_response = fallback_responses["default"]
            
            return jsonify({
                "response": fallback_response,
                "timestamp": datetime.utcnow().isoformat()
            })
        
    except Exception as e:
        print(f"[ERROR] Chat endpoint error: {e}")
        return jsonify({"error": "Sorry, I'm having trouble right now. Please try again."}), 500

@app.route("/api/upload", methods=["POST"])
@jwt_required()
def upload():
    try:
        user_id = get_jwt_identity()
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No selected file"}), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        meeting = Meeting(
            user_id=user_id,
            title=request.form.get('title', filename),
            filename=filename,
            language=request.form.get('language', 'en'),
            status='uploaded'
        )
        db.session.add(meeting)
        db.session.commit()
        
        return jsonify({"recording_id": meeting.id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/process/<int:meeting_id>", methods=["POST"])
@jwt_required()
def process_meeting(meeting_id):
    try:
        user_id = get_jwt_identity()
        meeting = Meeting.query.filter_by(id=meeting_id, user_id=user_id).first()
        
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        if meeting.status != 'uploaded':
            return jsonify({"error": f"Meeting already {meeting.status}"}), 400
        
        thread = threading.Thread(target=start_processing, args=(meeting_id,))
        thread.daemon = True
        thread.start()
        
        return jsonify({"message": "Processing started", "recording_id": meeting_id}), 202
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/processing-status/<int:meeting_id>", methods=["GET"])
@jwt_required()
def processing_status(meeting_id):
    try:
        user_id = get_jwt_identity()
        meeting = Meeting.query.filter_by(id=meeting_id, user_id=user_id).first()
        
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        try:
            steps = json.loads(meeting.processing_steps or '[]')
        except:
            steps = []
        
        if not steps:
            steps = [
                {"step": "transcription", "status": "pending", "timestamp": "", "error": None},
                {"step": "translation", "status": "pending", "timestamp": "", "error": None},
                {"step": "optimization", "status": "pending", "timestamp": "", "error": None},
                {"step": "ai_generation", "status": "pending", "timestamp": "", "error": None}
            ]
        
        return jsonify({
            "recording_id": meeting.id,
            "status": meeting.status,
            "processing_steps": steps,
            "current_step_progress": meeting.current_step_progress or 0
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/meetings", methods=["GET"])
@jwt_required()
def get_meetings():
    try:
        user_id = get_jwt_identity()
        limit = request.args.get('limit', 10000000, type=int)
        meetings = Meeting.query.filter_by(user_id=user_id).order_by(Meeting.upload_date.desc()).limit(limit).all()
        
        return jsonify({
            "meetings": [
                {
                    "id": m.id,
                    "title": m.title,
                    "filename": m.filename,
                    "upload_date": m.upload_date.isoformat(),
                    "status": m.status,
                    "has_transcription": m.has_transcription,
                    "has_notes": m.has_notes
                } for m in meetings
            ]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/meetings/<int:meeting_id>", methods=["DELETE"])
@jwt_required()
def delete_meeting(meeting_id):
    try:
        user_id = get_jwt_identity()
        meeting = Meeting.query.filter_by(id=meeting_id, user_id=user_id).first()
        
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], meeting.filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        
        db.session.delete(meeting)
        db.session.commit()
        
        return jsonify({"message": "Meeting deleted successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/meetings/<int:meeting_id>", methods=["GET"])
@jwt_required()
def get_meeting(meeting_id):
    try:
        user_id = get_jwt_identity()
        meeting = Meeting.query.filter_by(id=meeting_id, user_id=user_id).first()
        
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        return jsonify({
            "meeting": {
                "id": meeting.id,
                "title": meeting.title,
                "filename": meeting.filename,
                "upload_date": meeting.upload_date.isoformat(),
                "status": meeting.status,
                "transcription": json.loads(meeting.transcription or '{}'),
                "notes": json.loads(meeting.notes or '{}')
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/translate", methods=["POST"])
@jwt_required()
def translate_text():
    try:
        data = request.json
        text = data.get('text', '').strip()
        target_language = data.get('target_language', 'es')
        
        if not text:
            return jsonify({"error": "No text provided"}), 400
        
        language_names = {
            "af": "Afrikaans", "sq": "Albanian", "am": "Amharic", "ar": "Arabic", "hy": "Armenian",
            "az": "Azerbaijani", "eu": "Basque", "be": "Belarusian", "bn": "Bengali", "bs": "Bosnian",
            "bg": "Bulgarian", "ca": "Catalan", "ceb": "Cebuano", "ny": "Chichewa", "zh": "Chinese",
            "zh-cn": "Chinese (Simplified)", "zh-tw": "Chinese (Traditional)", "co": "Corsican",
            "hr": "Croatian", "cs": "Czech", "da": "Danish", "nl": "Dutch", "en": "English",
            "eo": "Esperanto", "et": "Estonian", "tl": "Filipino", "fi": "Finnish", "fr": "French",
            "fy": "Frisian", "gl": "Galician", "ka": "Georgian", "de": "German", "el": "Greek",
            "gu": "Gujarati", "ht": "Haitian Creole", "ha": "Hausa", "haw": "Hawaiian", "he": "Hebrew",
            "iw": "Hebrew", "hi": "Hindi", "hmn": "Hmong", "hu": "Hungarian", "is": "Icelandic",
            "ig": "Igbo", "id": "Indonesian", "ga": "Irish", "it": "Italian", "ja": "Japanese",
            "jw": "Javanese", "kn": "Kannada", "kk": "Kazakh", "km": "Khmer", "ko": "Korean",
            "ku": "Kurdish (Kurmanji)", "ky": "Kyrgyz", "lo": "Lao", "la": "Latin", "lv": "Latvian",
            "lt": "Lithuanian", "lb": "Luxembourgish", "mk": "Macedonian", "mg": "Malagasy",
            "ms": "Malay", "ml": "Malayalam", "mt": "Maltese", "mi": "Maori", "mr": "Marathi",
            "mn": "Mongolian", "my": "Myanmar (Burmese)", "ne": "Nepali", "no": "Norwegian",
            "or": "Odia", "ps": "Pashto", "fa": "Persian", "pl": "Polish", "pt": "Portuguese",
            "pa": "Punjabi", "ro": "Romanian", "ru": "Russian", "sm": "Samoan", "gd": "Scots Gaelic",
            "sr": "Serbian", "st": "Sesotho", "sn": "Shona", "sd": "Sindhi", "si": "Sinhala",
            "sk": "Slovak", "sl": "Slovenian", "so": "Somali", "es": "Spanish", "su": "Sundanese",
            "sw": "Swahili", "sv": "Swedish", "tg": "Tajik", "ta": "Tamil", "te": "Telugu",
            "th": "Thai", "tr": "Turkish", "uk": "Ukrainian", "ur": "Urdu", "ug": "Uyghur",
            "uz": "Uzbek", "vi": "Vietnamese", "cy": "Welsh", "xh": "Xhosa", "yi": "Yiddish",
            "yo": "Yoruba", "zu": "Zulu"
        }
        
        target_lang_name = language_names.get(target_language, "Spanish")
        
        try:
            prompt = f"""
You are a professional translator. Translate the given text to {target_lang_name}. Respond only with the translated text, no additional formatting or explanations.

Text to translate: {text}
"""
            response = call_gemini_api(prompt, model="gemini-1.5-flash")
            translated = response.text.strip() if response.text else ""
            
            if not translated:
                return jsonify({"error": "Translation failed: Empty response from API"}), 500
            
            return jsonify({"translated_text": translated})
            
        except Exception as e:
            print(f"[ERROR] Translation API error: {str(e)}")
            return jsonify({
                "error": "Translation service unavailable",
                "details": str(e),
                "suggestion": "Please check your GEMINI_API_KEY or network connection and try again."
            }), 500
        
    except Exception as e:
        print(f"[ERROR] Translate endpoint error: {str(e)}")
        return jsonify({"error": str(e)}), 500

def create_enhanced_pdf(meeting, filepath):
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        textColor='navy'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor='darkblue'
    )
    
    story.append(Paragraph(f"Meeting Notes: {meeting.title}", title_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph(f"<b>File:</b> {meeting.filename}", styles['Normal']))
    story.append(Paragraph(f"<b>Date:</b> {meeting.upload_date.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    try:
        notes = json.loads(meeting.notes or '{}')
        
        if notes.get('summary'):
            story.append(Paragraph("Executive Summary", heading_style))
            story.append(Paragraph(notes['summary'], styles['Normal']))
            story.append(Spacer(1, 15))
        
        if notes.get('key_points'):
            story.append(Paragraph("Key Discussion Points", heading_style))
            key_points = notes['key_points'] if isinstance(notes['key_points'], list) else json.loads(notes['key_points'] or '[]')
            for i, point in enumerate(key_points, 1):
                story.append(Paragraph(f"{i}. {point}", styles['Normal']))
            story.append(Spacer(1, 15))
        
        if notes.get('action_items'):
            story.append(Paragraph("Action Items", heading_style))
            action_items = notes['action_items'] if isinstance(notes['action_items'], list) else json.loads(notes['action_items'] or '[]')
            for i, item in enumerate(action_items, 1):
                story.append(Paragraph(f"• {item}", styles['Normal']))
            story.append(Spacer(1, 15))
        
        if notes.get('decisions'):
            story.append(Paragraph("Decisions Made", heading_style))
            decisions = notes['decisions'] if isinstance(notes['decisions'], list) else json.loads(notes['decisions'] or '[]')
            for decision in decisions:
                story.append(Paragraph(f"• {decision}", styles['Normal']))
            story.append(Spacer(1, 15))
        
        if notes.get('sentiment'):
            story.append(Paragraph("Overall Sentiment", heading_style))
            story.append(Paragraph(notes['sentiment'], styles['Normal']))
            story.append(Spacer(1, 15))
        
        transcription_data = json.loads(meeting.transcription or '{}')
        transcript_text = transcription_data.get('optimized') or transcription_data.get('translated') or transcription_data.get('raw')
        
        if transcript_text:
            story.append(PageBreak())
            story.append(Paragraph("Full Transcript", heading_style))
            story.append(Spacer(1, 10))
            
            max_chunk_size = 2000
            transcript_chunks = [transcript_text[i:i+max_chunk_size] for i in range(0, len(transcript_text), max_chunk_size)]
            
            for chunk in transcript_chunks:
                clean_chunk = chunk.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(clean_chunk, styles['Normal']))
                story.append(Spacer(1, 10))
        
    except Exception as e:
        story.append(Paragraph("Error: Could not parse meeting notes", styles['Normal']))
        print(f"[ERROR] PDF generation error: {e}")
    
    doc.build(story)

@app.route("/api/export/<int:id>/<string:format>", methods=["GET"])
@jwt_required()
def export(id, format):
    try:
        user_id = get_jwt_identity()
        meeting = Meeting.query.filter_by(id=id, user_id=user_id).first()
        
        if not meeting:
            return jsonify({"error": "Meeting not found"}), 404
        
        os.makedirs("outputs", exist_ok=True)
        
        if format == "word":
            filepath = f"outputs/meeting_notes_{id}.docx"
            doc = Document()
            doc.add_heading(f"Meeting Notes: {meeting.title}", 0)
            
            doc.add_paragraph(f"File: {meeting.filename}")
            doc.add_paragraph(f"Date: {meeting.upload_date.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")
            
            try:
                notes = json.loads(meeting.notes or '{}')
                
                if notes.get("summary"):
                    doc.add_heading("Executive Summary", level=1)
                    doc.add_paragraph(notes["summary"])
                
                if notes.get("key_points"):
                    doc.add_heading("Key Discussion Points", level=1)
                    key_points = notes["key_points"] if isinstance(notes["key_points"], list) else json.loads(notes["key_points"] or '[]')
                    for point in key_points:
                        doc.add_paragraph(point, style="List Bullet")
                
                if notes.get("action_items"):
                    doc.add_heading("Action Items", level=1)
                    action_items = notes["action_items"] if isinstance(notes["action_items"], list) else json.loads(notes["action_items"] or '[]')
                    for item in action_items:
                        doc.add_paragraph(item, style="List Number")
                
                if notes.get("decisions"):
                    doc.add_heading("Decisions Made", level=1)
                    decisions = notes["decisions"] if isinstance(notes["decisions"], list) else json.loads(notes["decisions"] or '[]')
                    for decision in decisions:
                        doc.add_paragraph(decision, style="List Bullet")
                
                if notes.get("sentiment"):
                    doc.add_heading("Overall Sentiment", level=1)
                    doc.add_paragraph(notes["sentiment"])
                
                transcription_data = json.loads(meeting.transcription or '{}')
                transcript_text = transcription_data.get('optimized') or transcription_data.get('translated') or transcription_data.get('raw')
                
                if transcript_text:
                    doc.add_heading("Full Transcript", level=1)
                    doc.add_paragraph(transcript_text)
                    
            except Exception as e:
                doc.add_paragraph(f"Error parsing notes: {str(e)}")
            
            doc.save(filepath)
            return send_file(filepath, as_attachment=True, download_name=f"meeting_notes_{id}.docx")
            
        elif format == "pdf":
            filepath = f"outputs/meeting_notes_{id}.pdf"
            create_enhanced_pdf(meeting, filepath)
            return send_file(filepath, as_attachment=True, download_name=f"meeting_notes_{id}.pdf")
        
        return jsonify({"error": "Invalid format"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/stats", methods=["GET"])
@jwt_required()
def stats():
    try:
        user_id = get_jwt_identity()
        meetings = Meeting.query.filter_by(user_id=user_id).all()
        
        total_uploads = len(meetings)
        total_words = sum(len(json.loads(m.notes or '{}').get("summary", "").split()) for m in meetings)
        
        today = datetime.utcnow().date()
        last_7_days = [(today - timedelta(days=i)).strftime("%a") for i in range(6, -1, -1)]
        uploads_by_day = Counter(m.upload_date.date().strftime("%a") for m in meetings)
        uploads_data = [uploads_by_day.get(day, 0) for day in last_7_days]
        
        return jsonify({
            "total_meetings": total_uploads,
            "completed_meetings": len([m for m in meetings if m.status == "completed"]),
            "total_words": total_words,
            "labels": last_7_days,
            "uploads": uploads_data
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    os.makedirs("uploads", exist_ok=True)
    os.makedirs("outputs", exist_ok=True)
    
    print("🚀 Starting Flask backend server...")
    print("📁 Upload folder:", app.config['UPLOAD_FOLDER'])
    print("🗄️ Database:", app.config['SQLALCHEMY_DATABASE_URI'])
    print("\n📧 Email Configuration:")
    print(f"   SMTP_USERNAME: {'✅ Set' if os.getenv('SMTP_USERNAME') else '❌ Not set'}")
    print(f"   SMTP_PASSWORD: {'✅ Set' if os.getenv('SMTP_PASSWORD') else '❌ Not set'}")
    print("\n   To enable email functionality:")
    print("   1. Create a .env file in your project root")
    print("   2. Add: SMTP_USERNAME=your-email@gmail.com")
    print("   3. Add: SMTP_PASSWORD=your-app-password")
    print("   4. For Gmail, use App Passwords: https://support.google.com/accounts/answer/185833")
    
    app.run(debug=True, host="0.0.0.0", port=5000)